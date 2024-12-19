from docx import Document
import os
from docx.shared import Inches

# Nama file untuk menyimpan data pesanan
FILENAME = 'pesanan-restoran-1.docx'

# Fungsi untuk memuat dokumen .docx
def load_document():
    if os.path.exists(FILENAME):
        return Document(FILENAME)
    else:
        doc = Document()
        doc.add_paragraph("Pesanan Restoran")
        doc.add_paragraph("=======================================")
        return doc

# Fungsi untuk menyimpan dokumen .docx
def save_document(doc):
    doc.save(FILENAME)

# Fungsi untuk menambah pesanan
def create_pesanan(nama, makanan, jumlah, gambar):
    doc = load_document()
    pesanan = doc.add_paragraph()
    pesanan.add_run(f"Nama Pelanggan: {nama}\n")
    pesanan.add_run(f"Makanan: {makanan}\n")
    pesanan.add_run(f"Jumlah: {jumlah}\n")
    pesanan.add_run(f"Image Path: {gambar}\n")

    try:
        # Menggunakan os.path.join untuk memastikan path relatif dikenali dengan benar
        image_path = os.path.join(os.getcwd(),gambar)
        pesanan.add_run().add_picture(image_path, width=Inches(2)) 
        pesanan.add_run("\n")
    except Exception as e:
        pesanan.add_run(f"Error saat menambahkan gambar: {str(e)}\n")

    pesanan.add_run(f"Status: diproses\n")
    pesanan.add_run("\n=======================================\n")
    save_document(doc)
    print("Pesanan berhasil ditambahkan!")

# Fungsi untuk menampilkan semua pesanan
def read_pesanan():
    doc = load_document()
    for para in doc.paragraphs:
        print(para.text)

# Fungsi untuk memperbarui status pesanan
def update_pesanan(nama, status_baru):
    doc = load_document()
    for para in doc.paragraphs:
        if nama in para.text and "Status" in para.text:
            old_status = para.text.split("Status: ")[1].strip()
            if old_status != "batal":
                # Hapus teks status sebelumnya tanpa menghapus gambar
                new_text = para.text.replace(f"Status: {old_status}", f"Status: {status_baru}")
                
                # Menghapus teks lama dan mengganti dengan teks baru
                para.clear()  # Menghapus seluruh isi paragraf
                para.add_run(new_text)  # Menambahkan teks baru dengan status yang diperbarui
                
                save_document(doc)
                print(f"Status pesanan {nama} berhasil diubah menjadi {status_baru}.")
                return
    print(f"Pesanan dengan nama {nama} tidak ditemukan atau status tidak dapat diubah.")


# Fungsi untuk menghapus pesanan yang berstatus batal
def delete_pesanan(nama):
    doc = load_document()
    paragraphs_to_delete = []
    
    # Iterasi paragraf dan simpan paragraf yang ingin dihapus
    for para in doc.paragraphs:
        if nama in para.text and "Status: batal" in para.text:
            paragraphs_to_delete.append(para)

    # Ganti paragraf yang ingin dihapus dengan teks kosong
    if paragraphs_to_delete:
        for para in paragraphs_to_delete:
            para.clear()  # Menghapus semua konten dari paragraf
        save_document(doc)
        print(f"Pesanan dengan nama {nama} berhasil dihapus.")
    else:
        print(f"Pesanan dengan nama {nama} tidak ditemukan atau tidak berstatus batal.")


# Fungsi untuk mencari pesanan berdasarkan nama pelanggan
def search_pesanan(nama):
    doc = load_document()
    found = False
    for para in doc.paragraphs:
        if nama.lower() in para.text.lower():
            print(para.text)
            found = True
    if not found:
        print(f"Pesanan dengan nama {nama} tidak ditemukan.")

# Menu utama
def main():
    while True:
        print("\n===== Manajemen Pesanan Restoran =====")
        print("1. Create Pesanan")
        print("2. Read Pesanan")
        print("3. Update Pesanan")
        print("4. Delete Pesanan")
        print("5. Search Pesanan")
        print("6. Exit")
        
        pilihan = input("Pilih opsi: ")
        
        if pilihan == "1":
            nama = input("Masukkan nama pelanggan: ")
            makanan = input("Masukkan nama makanan: ")
            jumlah = input("Masukkan jumlah makanan: ")
            gambar = input("Masukkan path file gambar: ")
            create_pesanan(nama, makanan, jumlah, gambar)
        
        elif pilihan == "2":
            print("\nMenampilkan Semua Pesanan: ")
            read_pesanan()
        
        elif pilihan == "3":
            nama = input("Masukkan nama pelanggan untuk mengubah status: ")
            status_baru = input("Masukkan status baru (proses/selesai/batal): ")
            update_pesanan(nama, status_baru)
        
        elif pilihan == "4":
            nama = input("Masukkan nama pelanggan yang ingin dihapus (status batal): ")
            delete_pesanan(nama)
        
        elif pilihan == "5":
            nama = input("Masukkan nama pelanggan yang ingin dicari: ")
            search_pesanan(nama)
        
        elif pilihan == "6":
            print("Keluar dari program.")
            break
        
        else:
            print("Opsi tidak valid, coba lagi.")

if __name__ == "__main__":
    main()
